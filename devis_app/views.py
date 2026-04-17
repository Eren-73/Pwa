from django.shortcuts import render, redirect, get_object_or_404
from .models import Devis, LigneDevis, Client, Categorie, Produit, PointVente, HistoriqueDevis, ResponsableCommercial
from .forms import DevisForm, LigneDevisForm,ClientForm,ProduitForm,CategorieForm
from django.forms import modelformset_factory
from django.urls import reverse
from .utils import generate_qr_code, normalize_phone_for_whatsapp
from django.http import JsonResponse
from django.contrib import messages
from django.template.loader import render_to_string
from docx import Document
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required  # 🔹 AJOUT
from django.contrib.auth.models import User
from django.utils import timezone
from django.db.models import Count
from django.db.models import Q
from django.db import transaction
from .models import Devis, Profile
import weasyprint
import openpyxl
import json
from datetime import datetime
from django.conf import settings
from .forms import CommercialCreateForm
from .models import ActionCommercial
import os
import io
from urllib.parse import quote
from collections import defaultdict
from django.core.mail import EmailMessage


def _get_profile(user):
    if not user.is_authenticated:
        return None
    try:
        return user.profile
    except Profile.DoesNotExist:
        return None


def _is_admin_user(user):
    if not user.is_authenticated:
        return False
    if user.is_superuser:
        return True
    profile = _get_profile(user)
    return bool(profile and profile.role == 'admin')


def _is_responsable_user(user):
    if not user.is_authenticated:
        return False
    if user.is_staff and not user.is_superuser:
        return True
    profile = _get_profile(user)
    return bool(profile and profile.role == 'responsable')


def _can_view_global(user):
    return _is_admin_user(user) or _is_responsable_user(user)


def _user_point_vente(user):
    profile = _get_profile(user)
    return profile.point_vente if profile else None


def _scoped_devis_queryset(user):
    base_qs = Devis.objects.select_related('client', 'point_vente', 'utilisateur', 'utilisateur__profile__point_vente')
    if not user.is_authenticated:
        return base_qs.none()

    if _can_view_global(user):
        return base_qs

    point_vente = _user_point_vente(user)
    if point_vente:
        return base_qs.filter(point_vente=point_vente)

    return base_qs.filter(utilisateur=user)


def _deletable_devis_queryset(user):
    """Devis que l'utilisateur a le droit de supprimer."""
    base_qs = Devis.objects.all()
    if not user.is_authenticated:
        return base_qs.none()

    # Admin/Responsable: suppression globale
    if _can_view_global(user):
        return base_qs

    # Commercial: suppression uniquement de ses propres devis
    return base_qs.filter(utilisateur=user)


def _can_modify_devis(user, devis):
    """Retourne True si l'utilisateur peut modifier ce devis."""
    if not user.is_authenticated:
        return False
    if _can_view_global(user):
        return True
    return devis.utilisateur_id == user.id


def _commercial_contact_for_devis(devis):
    """Retourne les informations de contact du commercial lié au devis."""
    if not devis.utilisateur:
        return {
            'name': 'PWA Energy Solution',
            'email': settings.DEFAULT_FROM_EMAIL,
            'phone': '',
        }

    profile = getattr(devis.utilisateur, 'profile', None)
    full_name = f"{devis.utilisateur.first_name} {devis.utilisateur.last_name}".strip() or devis.utilisateur.username
    phone = (getattr(profile, 'telephone', '') or '').strip()

    return {
        'name': full_name,
        'email': (devis.utilisateur.email or settings.DEFAULT_FROM_EMAIL).strip(),
        'phone': phone,
    }


def _whatsapp_phone(value):
    """Normalise un numéro de téléphone pour wa.me (format international)."""
    return normalize_phone_for_whatsapp(value)


def _is_ajax_request(request):
    return request.headers.get('X-Requested-With') == 'XMLHttpRequest'


def _scoped_clients_queryset(user):
    if not user.is_authenticated:
        return Client.objects.none()

    if _can_view_global(user):
        return Client.objects.all()

    # Les commerciaux doivent voir les clients existants, même avant qu'un devis leur soit lié.
    # On garde la restriction admin pour les actions sensibles (suppression, etc.).
    return Client.objects.all()

@login_required(login_url='login')
def creer_devis(request, slug=None):
    # Mode modification si slug est fourni
    devis_existant = None
    user_can_view_global = _can_view_global(request.user)
    user_point_vente = _user_point_vente(request.user)

    if slug:
        devis_existant = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
        if not _can_modify_devis(request.user, devis_existant):
            messages.error(request, "Vous n'avez pas le droit de modifier ce devis.")
            return redirect('dashboard')
    
    LigneDevisFormSet = modelformset_factory(
        LigneDevis, form=LigneDevisForm, extra=3 if slug else 0, can_delete=True
    )

    if request.method == 'POST':
        if slug:
            # Mode modification - Sauvegarder l'état avant modification
            lignes_data = []
            for ligne in devis_existant.lignes.all():
                lignes_data.append({
                    'produit': ligne.produit.nom if ligne.produit else '',
                    'designation': ligne.designation_display,
                    'quantite': float(ligne.quantite),
                    'unite': ligne.unite,
                    'pu': float(ligne.pu),
                    'pu_net': float(ligne.pu_net) if ligne.pu_net else float(ligne.pu),
                    'remise': float(ligne.remise),
                    'total_ht': float(ligne.total_ht),
                    'total_ttc': float(ligne.total_ttc) if ligne.total_ttc else float(ligne.total_ht)
                })
            
            donnees_avant = {
                'numero_devis': devis_existant.numero_devis,
                'date_emission': devis_existant.date_emission.isoformat(),
                'date_validite': devis_existant.date_validite.isoformat(),
                'date_proforma': devis_existant.date_proforma.isoformat() if devis_existant.date_proforma else devis_existant.date_emission.isoformat(),
                'total_ht': float(devis_existant.total_ht),
                'total_remise': float(devis_existant.total_remise) if devis_existant.total_remise else 0,
                'total_ht_remise': float(devis_existant.total_ht_remise) if devis_existant.total_ht_remise else float(devis_existant.total_ht),
                'total_tva': float(devis_existant.total_tva) if devis_existant.total_tva else 0,
                'total_ttc': float(devis_existant.total_ttc),
                'total_ttc_lettres': devis_existant.total_ttc_lettres if devis_existant.total_ttc_lettres else '',
                'regime_vente': devis_existant.regime_vente,
                'appliquer_tva': devis_existant.appliquer_tva,
                'detail_proposition': devis_existant.detail_proposition if devis_existant.detail_proposition else '',
                'pourcentage_acompte': float(devis_existant.pourcentage_acompte) if devis_existant.pourcentage_acompte else 60,
                'pourcentage_livraison': float(devis_existant.pourcentage_livraison) if devis_existant.pourcentage_livraison else 40,
                'client': {
                    'nom': devis_existant.client.nom if devis_existant.client else '',
                    'prenom': devis_existant.client.prenom if devis_existant.client else '',
                    'email': devis_existant.client.email if devis_existant.client else '',
                    'telephone': devis_existant.client.telephone if devis_existant.client else '',
                    'adresse': devis_existant.client.adresse if devis_existant.client else '',
                },
                'point_vente': {
                    'nom': devis_existant.point_vente.nom if devis_existant.point_vente else '',
                    'adresse': devis_existant.point_vente.adresse if devis_existant.point_vente else '',
                    'telephone': devis_existant.point_vente.telephone if devis_existant.point_vente else '',
                    'email': devis_existant.point_vente.email if devis_existant.point_vente else '',
                    'numero': devis_existant.point_vente.numero if devis_existant.point_vente else '',
                },
                'utilisateur': {
                    'first_name': devis_existant.utilisateur.first_name if devis_existant.utilisateur else '',
                    'last_name': devis_existant.utilisateur.last_name if devis_existant.utilisateur else '',
                    'email': devis_existant.utilisateur.email if devis_existant.utilisateur else '',
                    'telephone': devis_existant.utilisateur.profile.telephone if (devis_existant.utilisateur and hasattr(devis_existant.utilisateur, 'profile') and devis_existant.utilisateur.profile.telephone) else '',
                },
                'lignes': lignes_data
            }
            
            devis_form = DevisForm(request.POST, request.FILES, instance=devis_existant)
            formset = LigneDevisFormSet(request.POST, queryset=devis_existant.lignes.all())
        else:
            # Mode création
            devis_form = DevisForm(request.POST, request.FILES)
            formset = LigneDevisFormSet(request.POST, queryset=LigneDevis.objects.none())

        # Commercial: verrouiller le point de vente sur son site
        if not user_can_view_global and user_point_vente and 'point_vente' in devis_form.fields:
            devis_form.fields['point_vente'].queryset = PointVente.objects.filter(pk=user_point_vente.pk)

        if devis_form.is_valid() and formset.is_valid():
            try:
                with transaction.atomic():
                    devis = devis_form.save(commit=False)

                    if not slug:
                        # 🔹 Lier le devis à l'utilisateur connecté (commercial)
                        devis.utilisateur = request.user

                    # Les commerciaux restent dans leur point de vente
                    if not user_can_view_global and user_point_vente:
                        devis.point_vente = user_point_vente

                    # ✅ Vérifier si TVA doit être appliquée
                    apply_tva = request.POST.get("apply_tva", "yes")
                    devis.appliquer_tva = (apply_tva == "yes")

                    devis.save()

                    # Stock: préparer les quantités demandées sur le nouveau formset
                    requested_by_product = defaultdict(int)
                    kept_forms = []
                    for form in formset:
                        if not form.cleaned_data or form.cleaned_data.get('DELETE', False):
                            continue

                        produit = form.cleaned_data.get('produit')
                        quantite = int(form.cleaned_data.get('quantite') or 0)
                        if not produit:
                            continue
                        if quantite <= 0:
                            raise ValueError(f"Quantité invalide pour {produit.nom}.")

                        requested_by_product[produit.id] += quantite
                        kept_forms.append(form)

                    # Stock: récupérer les anciennes quantités en mode modification
                    previous_by_product = defaultdict(int)
                    if slug:
                        for old_line in devis_existant.lignes.select_related('produit'):
                            if old_line.produit_id:
                                previous_by_product[old_line.produit_id] += int(old_line.quantite)

                    involved_product_ids = set(requested_by_product.keys()) | set(previous_by_product.keys())
                    locked_products = {
                        p.id: p
                        for p in Produit.objects.select_for_update().filter(id__in=involved_product_ids)
                    }

                    # Validation stock avant toute écriture définitive
                    for product_id, requested_qty in requested_by_product.items():
                        product = locked_products.get(product_id)
                        if not product:
                            continue
                        available = int(product.stock) + int(previous_by_product.get(product_id, 0))
                        if requested_qty > available:
                            raise ValueError(
                                f"Stock insuffisant pour {product.nom}: demandé {requested_qty}, disponible {available}."
                            )

                    # En modification: restituer l'ancien stock puis supprimer les anciennes lignes
                    if slug:
                        for product_id, previous_qty in previous_by_product.items():
                            product = locked_products.get(product_id)
                            if product:
                                product.stock = int(product.stock) + int(previous_qty)
                                product.save(update_fields=['stock'])
                        devis.lignes.all().delete()

                    # Appliquer le nouveau stock (décrément)
                    for product_id, requested_qty in requested_by_product.items():
                        product = locked_products.get(product_id)
                        if product:
                            product.stock = int(product.stock) - int(requested_qty)
                            product.save(update_fields=['stock'])

                    # ✅ Enregistrer chaque ligne
                    for form in kept_forms:
                        ligne = form.save(commit=False)
                        ligne.devis = devis
                        ligne.save()
            except ValueError as stock_error:
                messages.error(request, str(stock_error))
                if slug:
                    devis_form = DevisForm(request.POST, request.FILES, instance=devis_existant)
                    formset = LigneDevisFormSet(request.POST, queryset=devis_existant.lignes.all())
                else:
                    devis_form = DevisForm(request.POST, request.FILES)
                    formset = LigneDevisFormSet(request.POST, queryset=LigneDevis.objects.none())
                return render(request, 'devis/creer_devis.html', {
                    'devis_form': devis_form,
                    'formset': formset,
                    'nombre_devis': Devis.objects.filter(utilisateur=request.user).count(),
                    'mode_modification': slug is not None,
                    'devis': devis_existant if slug else None,
                })

            if slug:
                # Enregistrer dans l'historique
                HistoriqueDevis.objects.create(
                    devis=devis,
                    utilisateur=request.user,
                    action='modification',
                    donnees_avant=donnees_avant,
                    commentaire=f"Modification par {request.user.username}"
                )
                messages.success(request, f"Devis n°{devis.numero_devis} modifié avec succès.")
            else:
                # ✅ Générer le lien public
                current_site_ip = "192.168.1.68"
                devis_url = f"http://{current_site_ip}:8000{reverse('devis_template', args=[devis.pk])}"

                # ✅ Générer le QR code
                print("Lien dans le QR :", devis_url)
                devis.qr_code.save(
                    f"qr_{devis.slug}.png",
                    generate_qr_code(devis.numero_devis),
                    save=True
                )
                # Ajout d'un message de succès après création
                messages.success(request, f"Proposition commerciale n°{devis.numero_devis} créée avec succès.")

            return redirect('devis_template', slug=devis.slug)
        else:
            messages.error(request, "Impossible d'enregistrer le devis. Vérifie les champs du formulaire et les lignes produits.")

    else:
        if slug:
            # Mode modification - pré-remplir avec données existantes
            devis_form = DevisForm(instance=devis_existant)
            formset = LigneDevisFormSet(queryset=devis_existant.lignes.all())
        else:
            # Mode création
            devis_form = DevisForm()
            formset = LigneDevisFormSet(queryset=LigneDevis.objects.none())

        # Commercial: verrouiller le point de vente sur son site
        if not user_can_view_global and user_point_vente and 'point_vente' in devis_form.fields:
            devis_form.fields['point_vente'].queryset = PointVente.objects.filter(pk=user_point_vente.pk)
            devis_form.fields['point_vente'].initial = user_point_vente

    # 🔹 Nombre total de devis créés par l'utilisateur connecté
    nombre_devis = Devis.objects.filter(utilisateur=request.user).count()

    return render(request, 'devis/creer_devis.html', {
        'devis_form': devis_form,
        'formset': formset,
        'nombre_devis': nombre_devis,
        'mode_modification': slug is not None,
        'devis': devis_existant if slug else None,
    })

def devis_template(request, slug):
    devis = get_object_or_404(Devis, slug=slug)
    lignes = devis.lignes.all()
    can_modify_devis = _can_modify_devis(request.user, devis)
    commercial_contact = _commercial_contact_for_devis(devis)
    whatsapp_url = ''

    client_phone = _whatsapp_phone(getattr(devis.client, 'telephone', ''))
    if client_phone:
        pdf_url = request.build_absolute_uri(reverse('export_pdf', args=[devis.slug]))
        message = (
            f"Bonjour {devis.client.prenom} {devis.client.nom}, "
            f"voici votre devis N° {devis.numero_devis} (Total TTC: {devis.total_ttc} CFA). "
            f"Téléchargez le PDF ici: {pdf_url} "
            f"Contact commercial: {commercial_contact['name']}"
        )
        if commercial_contact['phone']:
            message += f" - WhatsApp: {commercial_contact['phone']}"
        whatsapp_url = f"https://api.whatsapp.com/send/?phone={client_phone}&text={quote(message)}&type=phone_number&app_absent=0"

    return render(request, 'devis/devis_template.html', {
        'devis': devis,
        'lignes': lignes,
        'qr_code_url': request.build_absolute_uri(devis.qr_code.url) if devis.qr_code else None,
        'can_modify_devis': can_modify_devis,
        'commercial_contact': commercial_contact,
        'whatsapp_url': whatsapp_url,
        'whatsapp_client_phone': client_phone,
    })

@login_required(login_url='login')
def dashboard(request):
    # Filtrer par date si des paramètres GET sont passés
    date_debut = request.GET.get('date_debut')
    date_fin = request.GET.get('date_fin')
    query = (request.GET.get('q') or '').strip()

    devis_list = _scoped_devis_queryset(request.user).order_by('-date_emission')

    if date_debut and date_fin:
        devis_list = devis_list.filter(
            date_emission__range=[date_debut, date_fin]
        )

    if query:
        devis_list = devis_list.filter(
            Q(numero_devis__icontains=query)
            | Q(client__nom__icontains=query)
            | Q(client__prenom__icontains=query)
            | Q(client__email__icontains=query)
            | Q(detail_proposition__icontains=query)
            | Q(point_vente__nom__icontains=query)
        )

    deletable_ids = set(_deletable_devis_queryset(request.user).values_list('id', flat=True))

    # Site affiché: point de vente du devis, sinon site du commercial créateur.
    for devis in devis_list:
        resolved_site = devis.point_vente
        if not resolved_site and devis.utilisateur and hasattr(devis.utilisateur, 'profile'):
            resolved_site = devis.utilisateur.profile.point_vente
        devis.resolved_site = resolved_site
        devis.can_delete = devis.id in deletable_ids

    role_label = 'Commercial'
    if _is_admin_user(request.user):
        role_label = 'Admin'
    elif _is_responsable_user(request.user):
        role_label = 'Responsable'

    return render(request, 'dashboard.html', {
        'devis_list': devis_list,
        'q': query,
        'can_view_global': _can_view_global(request.user),
        'is_admin_user': _is_admin_user(request.user),
        'can_delete_devis': request.user.is_authenticated,
        'role_label': role_label,
    })




@login_required(login_url='login')
def liste_clients(request):
    query = (request.GET.get('q') or '').strip()
    clients = _scoped_clients_queryset(request.user).order_by('-date_creation')

    if query:
        clients = clients.filter(
            Q(nom__icontains=query)
            | Q(prenom__icontains=query)
            | Q(email__icontains=query)
            | Q(telephone__icontains=query)
            | Q(devis__numero_devis__icontains=query)
        ).distinct()

    # Le client n'est pas assigné à un point de vente dans la base.
    # On l'affiche selon le point de vente du commercial qui l'a enregistré.
    client_slugs = [client.slug for client in clients if client.slug]
    slug_set = set(client_slugs)
    creator_by_slug = {}
    if slug_set:
        creation_logs = (
            ActionCommercial.objects
            .filter(action__startswith='Création du client ')
            .select_related('user__profile__point_vente')
            .order_by('-date_action')
        )
        for log in creation_logs:
            slug = log.action.replace('Création du client ', '', 1).strip()
            if slug in slug_set and slug not in creator_by_slug:
                creator_by_slug[slug] = log.user

    for client in clients:
        creator = creator_by_slug.get(client.slug)
        if creator and hasattr(creator, 'profile') and creator.profile.point_vente:
            client.point_vente_label = creator.profile.point_vente.nom
        else:
            client.point_vente_label = 'Non assigné'

    return render(request, 'clients/clients.html', {
        'clients': clients,
        'q': query,
        'can_view_global': _can_view_global(request.user),
        'can_delete_clients': _is_admin_user(request.user),
    })


@login_required(login_url='login')
def ajouter_client(request):
    if request.method == 'POST':
        form = ClientForm(request.POST)
        if form.is_valid():
            client = form.save()
            ActionCommercial.objects.create(
                user=request.user,
                action=f"Création du client {client.slug}"
            )
            messages.success(request, "Client ajouté avec succès.")
            return redirect('liste_clients')
    else:
        form = ClientForm()

    if request.method == 'POST' and form.errors:
        messages.error(request, "Impossible d'ajouter le client. Vérifie les champs saisis.")

    return render(request, 'clients/ajouter_client.html', {'form': form})

@login_required(login_url='login')
def modifier_client(request, slug):
    client = get_object_or_404(_scoped_clients_queryset(request.user), slug=slug)
    if request.method == 'POST':
        form = ClientForm(request.POST, instance=client)
        if form.is_valid():
            form.save()
            messages.success(request, "Client modifié avec succès.")
            return redirect('liste_clients')
        messages.error(request, "Impossible de modifier le client. Vérifie les champs saisis.")
    else:
        form = ClientForm(instance=client)

    return render(request, 'clients/modifier_client.html', {'form': form, 'client': client})


@login_required(login_url='login')
def supprimer_client(request, slug):
    if not _is_admin_user(request.user):
        if _is_ajax_request(request):
            return JsonResponse({"success": False, "error": "Accès non autorisé."}, status=403)
        messages.error(request, "Accès non autorisé.")
        return redirect('liste_clients')

    client = get_object_or_404(_scoped_clients_queryset(request.user), slug=slug)
    if request.method == 'POST':
        client_nom = str(client)
        client.delete()
        if _is_ajax_request(request):
            return JsonResponse({"success": True, "deleted": client_nom})
        return redirect('/clients/?deleted=1')
    return render(request, 'supprimer_client.html', {'client': client})

@login_required(login_url='login')
def devis_par_client(request, slug):
    client = get_object_or_404(_scoped_clients_queryset(request.user), slug=slug)
    devis_list = _scoped_devis_queryset(request.user).filter(client=client)
    return render(request, 'clients/devis_par_client.html', {
        'client': client,
        'devis_list': devis_list
    })

@login_required(login_url='login')
def detail_devis(request, slug):
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    # Récupère les lignes associées au devis pour l'affichage
    lignes = devis.lignes.all()

    return render(request, "devis/devis_template.html", {
        "devis": devis,
        "lignes": lignes,
    })


@login_required(login_url='login')
def modifier_devis(request, slug):
    """Modifier un devis existant en conservant le numéro"""
    from .models import HistoriqueDevis
    import json
    from django.forms import model_to_dict
    
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    LigneDevisFormSet = modelformset_factory(
        LigneDevis, form=LigneDevisForm, extra=3, can_delete=True
    )
    
    if request.method == 'POST':
        # Sauvegarder l'état avant modification (convertir Decimal en float pour JSON)
        lignes_data = []
        for ligne in devis.lignes.all():
            lignes_data.append({
                'produit': ligne.produit.nom if ligne.produit else '',
                'designation': ligne.designation_display,
                'quantite': float(ligne.quantite),
                'unite': ligne.unite,
                'pu': float(ligne.pu),
                'pu_net': float(ligne.pu_net) if ligne.pu_net else float(ligne.pu),
                'remise': float(ligne.remise),
                'total_ht': float(ligne.total_ht),
                'total_ttc': float(ligne.total_ttc) if ligne.total_ttc else float(ligne.total_ht)
            })
        
        donnees_avant = {
            'numero_devis': devis.numero_devis,
            'date_emission': devis.date_emission.isoformat(),
            'date_validite': devis.date_validite.isoformat(),
            'date_proforma': devis.date_proforma.isoformat() if devis.date_proforma else devis.date_emission.isoformat(),
            'total_ht': float(devis.total_ht),
            'total_remise': float(devis.total_remise) if devis.total_remise else 0,
            'total_ht_remise': float(devis.total_ht_remise) if devis.total_ht_remise else float(devis.total_ht),
            'total_tva': float(devis.total_tva) if devis.total_tva else 0,
            'total_ttc': float(devis.total_ttc),
            'total_ttc_lettres': devis.total_ttc_lettres if devis.total_ttc_lettres else '',
            'regime_vente': devis.regime_vente,
            'appliquer_tva': devis.appliquer_tva,
            'detail_proposition': devis.detail_proposition if devis.detail_proposition else '',
            'pourcentage_acompte': float(devis.pourcentage_acompte) if devis.pourcentage_acompte else 60,
            'pourcentage_livraison': float(devis.pourcentage_livraison) if devis.pourcentage_livraison else 40,
            'client': {
                'nom': devis.client.nom if devis.client else '',
                'prenom': devis.client.prenom if devis.client else '',
                'email': devis.client.email if devis.client else '',
                'telephone': devis.client.telephone if devis.client else '',
                'adresse': devis.client.adresse if devis.client else '',
            },
            'point_vente': {
                'nom': devis.point_vente.nom if devis.point_vente else '',
                'adresse': devis.point_vente.adresse if devis.point_vente else '',
                'telephone': devis.point_vente.telephone if devis.point_vente else '',
                'email': devis.point_vente.email if devis.point_vente else '',
                'numero': devis.point_vente.numero if devis.point_vente else '',
            },
            'utilisateur': {
                'first_name': devis.utilisateur.first_name if devis.utilisateur else '',
                'last_name': devis.utilisateur.last_name if devis.utilisateur else '',
                'email': devis.utilisateur.email if devis.utilisateur else '',
                'telephone': devis.utilisateur.profile.telephone if (devis.utilisateur and hasattr(devis.utilisateur, 'profile') and devis.utilisateur.profile.telephone) else '',
            },
            'lignes': lignes_data
        }
        
        devis_form = DevisForm(request.POST, instance=devis)
        formset = LigneDevisFormSet(request.POST, queryset=devis.lignes.all())
        
        if devis_form.is_valid() and formset.is_valid():
            # Sauvegarder le devis avec mise à jour de la date uniquement
            devis = devis_form.save(commit=False)

            # Commercial: verrouiller le point de vente sur son site
            if not _can_view_global(request.user):
                user_point_vente = _user_point_vente(request.user)
                if user_point_vente:
                    devis.point_vente = user_point_vente
            
            # Vérifier la TVA
            apply_tva = request.POST.get("apply_tva", "yes")
            devis.appliquer_tva = (apply_tva == "yes")
            
            devis.save()
            
            # Supprimer les anciennes lignes et créer les nouvelles
            devis.lignes.all().delete()
            for form in formset:
                if form.cleaned_data and not form.cleaned_data.get('DELETE', False):
                    if not form.cleaned_data.get('produit'):
                        continue
                    ligne = form.save(commit=False)
                    ligne.devis = devis
                    ligne.save()
            
            # Enregistrer dans l'historique
            HistoriqueDevis.objects.create(
                devis=devis,
                utilisateur=request.user,
                action='modification',
                donnees_avant=donnees_avant,
                commentaire=f"Modification par {request.user.username}"
            )
            
            messages.success(request, f"Devis n°{devis.numero_devis} modifié avec succès.")
            return redirect('devis_template', slug=devis.slug)
        else:
            # Afficher les erreurs
            if not devis_form.is_valid():
                messages.error(request, f"Erreurs dans le formulaire : {devis_form.errors}")
            if not formset.is_valid():
                messages.error(request, f"Erreurs dans les lignes : {formset.errors}")
    else:
        devis_form = DevisForm(instance=devis)
        formset = LigneDevisFormSet(queryset=devis.lignes.all())
    
    return render(request, 'devis/modifier_devis.html', {
        'devis_form': devis_form,
        'formset': formset,
        'devis': devis,
    })


@login_required(login_url='login')
def historique_devis(request, slug):
    """Afficher l'historique des modifications d'un devis"""
    from .models import HistoriqueDevis
    
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    historique = HistoriqueDevis.objects.filter(devis=devis).order_by('-date_modification')
    
    return render(request, 'devis/historique_devis.html', {
        'devis': devis,
        'historique': historique,
    })


@login_required(login_url='login')
def voir_version_historique(request, historique_id):
    """
    Affiche une version historique du devis comme une facture complète
    """
    historique = get_object_or_404(HistoriqueDevis, id=historique_id, devis__in=_scoped_devis_queryset(request.user))
    devis_actuel = historique.devis
    
    # Créer un objet devis temporaire avec les données historiques
    if historique.donnees_avant:
        # On va passer les données historiques au template
        context = {
            'devis': devis_actuel,
            'historique': historique,
            'donnees_avant': historique.donnees_avant,
            'est_version_historique': True,
        }
        return render(request, 'devis/version_historique.html', context)
    else:
        messages.warning(request, "Aucune donnée historique disponible pour cette version.")
        return redirect('historique_devis', slug=devis_actuel.slug)


@login_required(login_url='login')
def supprimer_devis_selectionnes(request):
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Méthode non autorisée"}, status=405)

    content_type = request.headers.get('Content-Type', '')
    wants_json = (
        'application/json' in content_type
        or request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    )

    try:
        raw_ids = []

        if 'application/json' in content_type:
            try:
                data = json.loads((request.body or b'{}').decode('utf-8'))
            except json.JSONDecodeError:
                if wants_json:
                    return JsonResponse({"success": False, "error": "Payload JSON invalide."}, status=400)
                messages.error(request, "Requête de suppression invalide.")
                return redirect('dashboard')
            raw_ids = data.get('devis_ids', [])
        else:
            # Fallback pour formulaires classiques
            raw_ids = request.POST.getlist('devis_ids[]') or request.POST.getlist('devis_ids')

        devis_ids = []
        for value in raw_ids:
            try:
                devis_ids.append(int(value))
            except (TypeError, ValueError):
                continue

        # Déduplique pour éviter les suppressions redondantes
        devis_ids = list(set(devis_ids))

        if not devis_ids:
            if wants_json:
                return JsonResponse({"success": False, "error": "Aucun devis sélectionné."}, status=400)
            messages.warning(request, "Aucun devis sélectionné.")
            return redirect('dashboard')

        deleted_count, _ = _deletable_devis_queryset(request.user).filter(id__in=devis_ids).delete()
        if deleted_count == 0:
            if wants_json:
                return JsonResponse({"success": False, "error": "Aucun devis supprimé (droits insuffisants)."}, status=403)
            messages.error(request, "Aucun devis supprimé (droits insuffisants).")
            return redirect('dashboard')

        if wants_json:
            return JsonResponse({"success": True, "deleted_count": deleted_count})

        messages.success(request, f"{deleted_count} devis supprimé(s) avec succès.")
        return redirect('dashboard')
    except Exception as e:
        if wants_json:
            return JsonResponse({"success": False, "error": str(e)}, status=400)
        messages.error(request, f"Erreur pendant la suppression: {str(e)}")
        return redirect('dashboard')


@login_required(login_url='login')
def envoyer_devis_par_email(request, slug):
    """
    Envoie le devis par email au client avec le PDF en pièce jointe
    """
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    lignes = devis.lignes.all()
    commercial_contact = _commercial_contact_for_devis(devis)
    
    # Vérifier si le client a un email
    if not devis.client or not devis.client.email:
        messages.error(request, "Le client n'a pas d'adresse email.")
        return redirect('devis_template', slug=slug)
    
    # Générer le PDF
    html_string = render_to_string('devis/devis_template.html', {
        'devis': devis,
        'lignes': lignes,
        'qr_code_url': request.build_absolute_uri(devis.qr_code.url) if devis.qr_code else None,
        'is_pdf': True,
    })
    
    html = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri())
    pdf = html.write_pdf()
    
    # Préparer l'email
    sujet = f"Devis N° {devis.numero_devis} - PWA Energy Solution"
    message = f"""Bonjour {devis.client.prenom} {devis.client.nom},

Veuillez trouver ci-joint votre devis N° {devis.numero_devis} d'un montant de {devis.total_ttc} CFA.

Détails du devis:
- Date d'émission: {devis.date_emission.strftime('%d/%m/%Y')}
- Date de validité: {devis.date_validite.strftime('%d/%m/%Y')}
- Montant Total TTC: {devis.total_ttc} CFA

N'hésitez pas à nous contacter pour toute question.

Commercial: {commercial_contact['name']}
Email: {commercial_contact['email']}
Téléphone/WhatsApp: {commercial_contact['phone'] or 'Non renseigné'}

Cordialement,
PWA Energy Solution
"""
    
    # Créer et envoyer l'email
    email = EmailMessage(
        subject=sujet,
        body=message,
        # SMTP gratuit: utiliser un expéditeur unique vérifié, puis répondre au commercial.
        from_email=settings.DEFAULT_FROM_EMAIL,
        to=[devis.client.email],
        reply_to=[commercial_contact['email']] if commercial_contact['email'] else [],
    )
    
    # Attacher le PDF
    email.attach(f'devis_{devis.numero_devis}.pdf', pdf, 'application/pdf')
    
    try:
        email.send()
        messages.success(request, f"Devis envoyé avec succès à {devis.client.email}")
    except Exception as e:
        messages.error(request, f"Erreur lors de l'envoi de l'email: {str(e)}")
    
    return redirect('devis_template', slug=slug)


@login_required(login_url='login')
def export_devis_excel(request, slug):
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    lignes = devis.lignes.all()

    # Create workbook and worksheet
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Devis {devis.numero_devis}"

    # Styles
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, NamedStyle
    from openpyxl.styles.numbers import FORMAT_NUMBER_COMMA_SEPARATED1

    # Fonts
    title_font = Font(name='Arial', bold=True, size=14)
    header_font = Font(name='Arial', bold=True, size=11)
    normal_font = Font(name='Arial', size=10)
    bold = Font(name='Arial', bold=True, size=10)

    # Alignments
    right_align = Alignment(horizontal='right', vertical='center')
    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')

    # Borders
    thin = Side(border_style='thin', color='000000')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    top_border = Border(top=thin)

    # Fill colors
    light_gray = 'E6E6E6'
    lighter_gray = 'F2F2F2'

    # Money format
    money_format = '#,##0.00 "XOF"'

    # Column widths optimized for content
    widths = {
        'A': 45,  # Designation
        'B': 8,   # Qté
        'C': 8,   # Unité
        'D': 15,  # Prix unitaire
        'E': 10,  # Remise
        'F': 15,  # P.U Net
        'G': 15,  # Total
    }
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    # Row heights
    ws.row_dimensions[1].height = 30  # Logo row
    for row in range(2, 30):  # Adjust other rows
        ws.row_dimensions[row].height = 20

    # Insert logo if available
    try:
        logo_path = os.path.join(settings.STATIC_ROOT or settings.BASE_DIR, 'images', 'logo.png')
        if os.path.exists(logo_path):
            img = openpyxl.drawing.image.Image(logo_path)
            img.width = 200
            img.height = 90
            ws.add_image(img, 'A1')
    except Exception:
        pass

    # Header texts
    ws['D1'] = 'PROPOSITION COMMERCIALE'
    ws['D2'] = 'PROFORMA'
    for cell in ['D1', 'D2']:
        ws[cell].font = title_font
        ws[cell].alignment = right_align

    # Devis meta with styling
    ws['A4'] = f'Devis N° {devis.numero_devis}'
    ws['A5'] = f'Date: {devis.date_emission.strftime("%d/%m/%Y")}'
    for cell in ['A4', 'A5']:
        ws[cell].font = bold
        ws[cell].border = top_border

    # Client block with consistent styling
    row = 7
    ws[f'A{row}'] = 'CLIENT'
    ws[f'A{row}'].font = header_font
    ws[f'A{row}'].fill = PatternFill(start_color=light_gray, end_color=light_gray, fill_type='solid')
    row += 1

    client_data = [
        ('Nom', f"{devis.client.nom} {devis.client.prenom}"),
        ('Email', devis.client.email),
        ('Adresse', devis.client.adresse or ''),
    ]
    for label, value in client_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].font = bold
        ws[f'B{row}'].font = normal_font
        ws[f'A{row}'].alignment = left_align
        ws[f'B{row}'].alignment = left_align
        row += 1

    # Commercial block with matching style
    crow = 7
    ws[f'D{crow}'] = 'CONSEILLER COMMERCIAL'
    ws[f'D{crow}'].font = header_font
    ws[f'D{crow}'].fill = PatternFill(start_color=light_gray, end_color=light_gray, fill_type='solid')
    crow += 1

    if devis.utilisateur:
        commercial_data = [
            ('Nom', f"{devis.utilisateur.first_name} {devis.utilisateur.last_name}"),
            ('Email', devis.utilisateur.email),
            ('Point de vente', devis.point_vente.nom if devis.point_vente else 'Non assigné'),
        ]
        for label, value in commercial_data:
            ws[f'D{crow}'] = label
            ws[f'E{crow}'] = value
            ws[f'D{crow}'].font = bold
            ws[f'E{crow}'].font = normal_font
            ws[f'D{crow}'].alignment = left_align
            ws[f'E{crow}'].alignment = left_align
            crow += 1
    else:
        ws[f'D{crow}'] = 'Commercial non assigné'
        ws[f'D{crow}'].font = normal_font

    # Products table with enhanced styling
    table_row = 12
    headers = ['Désignation', 'Qté', 'Unité', 'Prix unitaire', 'Remise', 'P.U Net', 'Total']
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=table_row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = PatternFill(start_color=light_gray, end_color=light_gray, fill_type='solid')
        cell.alignment = center_align
        cell.border = border

    # Format product rows with proper number formatting
    for ligne in lignes:
        table_row += 1
        # Row data with proper types
        values = [
            (ligne.designation_display, 'text', left_align),
            (ligne.quantite, 'number', right_align),
            (ligne.unite, 'text', center_align),
            (ligne.pu, 'money', right_align),
            (f"{ligne.remise}%", 'text', right_align),
            (ligne.pu_net, 'money', right_align),
            (ligne.total_ttc, 'money', right_align),
        ]
        
        for col_idx, (val, type_, align) in enumerate(values, 1):
            cell = ws.cell(row=table_row, column=col_idx, value=val)
            cell.font = normal_font
            cell.alignment = align
            cell.border = border
            
            if type_ == 'money':
                cell.number_format = money_format

    # Totals area with enhanced styling
    totals_start = table_row + 2
    totals = [
        ('Total HT', devis.total_ht),
        ('Remise', devis.total_remise),
        ('Total HT Remise', devis.total_ht_remise),
        ('TVA (18%)', devis.total_tva),
        ('Total TTC', devis.total_ttc),
    ]
    
    # Style the totals section
    tr = totals_start
    for label, value in totals:
        # Label in column F
        label_cell = ws.cell(row=tr, column=6, value=label)
        label_cell.font = bold
        label_cell.alignment = right_align
        if label == 'Total TTC':
            label_cell.fill = PatternFill(start_color=lighter_gray, end_color=lighter_gray, fill_type='solid')
        
        # Value in column G
        val_cell = ws.cell(row=tr, column=7, value=value)
        val_cell.font = bold
        val_cell.number_format = money_format
        if label == 'Total TTC':
            val_cell.fill = PatternFill(start_color=lighter_gray, end_color=lighter_gray, fill_type='solid')
        
        # Add top border for Total TTC
        if label == 'Total TTC':
            label_cell.border = Border(top=Side(border_style='double'))
            val_cell.border = Border(top=Side(border_style='double'))
        
        tr += 1

    # Total en lettres with proper styling
    words_cell = ws.cell(row=tr + 1, column=1, value=f"Total en lettres: {devis.total_ttc_lettres}")
    words_cell.font = bold
    words_cell.alignment = left_align
    ws.merge_cells(start_row=tr + 1, start_column=1, end_row=tr + 1, end_column=5)

    # Response
    filename = f"devis_{devis.numero_devis}_{datetime.now().strftime('%Y%m%d')}.xlsx"
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response


@login_required(login_url='login')
def liste_materiels(request):
    """Affiche la liste des matériels / produits.

    Cette vue est référencée par `devis_app/urls.py` (path 'materiels/').
    Elle renvoie les catégories (et leurs produits) vers le template `liste_materiels.html`.
    Si le template n'existe pas encore, créer un fichier minimal dans `devis_app/templates/`.
    """
    categories = Categorie.objects.all().prefetch_related('produit_set')
    produits = Produit.objects.all()
    return render(request, 'produits/liste_materiels.html', {
        'categories': categories,
        'produits': produits,
    })


@login_required(login_url='login')
def ajouter_produit(request):
    """Stub view to add a product. Implements minimal POST handling using ProduitForm.

    If you want full behaviour, we can expand this to match your admin UX.
    """
    if request.method == 'POST':
        form = ProduitForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('liste_materiels')
    else:
        form = ProduitForm()
    return render(request, 'produits/modifier_produit.html', {'form': form})


@login_required(login_url='login')
def modifier_produit(request, produit_id):
    produit = get_object_or_404(Produit, id=produit_id)
    if request.method == 'POST':
        form = ProduitForm(request.POST, instance=produit)
        if form.is_valid():
            form.save()
            return redirect('liste_materiels')
    else:
        form = ProduitForm(instance=produit)
    return render(request, 'produits/modifier_produit.html', {'form': form, 'produit': produit})


@login_required(login_url='login')
def supprimer_produit(request, produit_id):
    produit = get_object_or_404(Produit, id=produit_id)
    if request.method == 'POST':
        produit_nom = produit.nom
        produit.delete()
        if _is_ajax_request(request):
            return JsonResponse({"success": True, "deleted": produit_nom})
        return redirect('liste_materiels')
    return render(request, 'supprimer_produit.html', {'produit': produit})


@login_required(login_url='login')
def ajouter_categorie(request):
    if request.method == 'POST':
        form = CategorieForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('liste_materiels')
    else:
        form = CategorieForm()
    return render(request, 'modifier_categorie.html', {'form': form})


@login_required(login_url='login')
def export_pdf(request, slug):
    """Exporte le devis en PDF avec le même template que l'écran."""
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    lignes = devis.lignes.all()

    commercial_contact = _commercial_contact_for_devis(devis)
    context = {
        'devis': devis,
        'lignes': lignes,
        'qr_code_url': request.build_absolute_uri(devis.qr_code.url) if devis.qr_code else None,
        'is_pdf': True,
        'commercial_contact': commercial_contact,
    }

    html = render_to_string('devis/devis_template.html', context, request=request)
    
    try:
        pdf = weasyprint.HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
        
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="devis_{devis.numero_devis}.pdf"'
        return response
    except Exception as e:
        print(f"PDF Generation Error: {str(e)}")  # Debug log
        return HttpResponse(html)  # Fallback to HTML for debugging


@login_required(login_url='login')
def export_devis_word(request, slug):
    """Minimal Word export: creates a simple .docx with basic devis info.

    This is intentionally lightweight to ensure the view exists and works.
    We can expand formatting later to match your desired layout exactly.
    """
    devis = get_object_or_404(_scoped_devis_queryset(request.user), slug=slug)
    lignes = devis.lignes.all()

    doc = Document()
    doc.add_heading(f'Devis N° {devis.numero_devis}', level=1)
    doc.add_paragraph(f'Date: {devis.date_emission.strftime("%d/%m/%Y")}')
    doc.add_paragraph(f'Client: {devis.client.nom} {devis.client.prenom}')
    doc.add_paragraph('')
    doc.add_paragraph('Lignes :')
    for l in lignes:
        doc.add_paragraph(f'- {l.designation_display} x{l.quantite} : {l.total_ttc}')

    # prepare response
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"devis_{devis.numero_devis}_{datetime.now().strftime('%Y%m%d')}.docx"
    response = HttpResponse(bio.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required(login_url='login')
def liste_point_ventes(request):
    """
    Liste tous les points de vente.
    """
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    points = PointVente.objects.all().order_by('nom')
    return render(request, 'points_vente/point_vente_list.html', {'points': points})


@login_required(login_url='login')
def ajouter_point_vente(request):
    """
    Ajouter un nouveau point de vente.
    """
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    from .forms import PointVenteForm
    
    if request.method == 'POST':
        form = PointVenteForm(request.POST)
        if form.is_valid():
            point_vente = form.save()
            messages.success(request, f"Point de vente '{point_vente.nom}' ajouté avec succès!")
            return redirect('liste_point_ventes')
    else:
        form = PointVenteForm()
    
    return render(request, 'points_vente/ajouter_point_vente.html', {'form': form})


@login_required(login_url='login')
def modifier_point_vente(request, pk):
    """
    Modifier un point de vente existant.
    """
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    from .forms import PointVenteForm
    
    point_vente = get_object_or_404(PointVente, pk=pk)
    
    if request.method == 'POST':
        form = PointVenteForm(request.POST, instance=point_vente)
        if form.is_valid():
            form.save()
            messages.success(request, f"Point de vente '{point_vente.nom}' modifié avec succès!")
            return redirect('liste_point_ventes')
    else:
        form = PointVenteForm(instance=point_vente)
    
    return render(request, 'points_vente/modifier_point_vente.html', {
        'form': form,
        'point_vente': point_vente
    })


@login_required(login_url='login')
def supprimer_point_vente(request, pk):
    """
    Supprimer un point de vente.
    """
    if not _is_admin_user(request.user):
        if _is_ajax_request(request):
            return JsonResponse({"success": False, "error": "Accès non autorisé."}, status=403)
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    point_vente = get_object_or_404(PointVente, pk=pk)
    
    if request.method == 'POST':
        nom = point_vente.nom
        point_vente.delete()
        if _is_ajax_request(request):
            return JsonResponse({"success": True, "deleted": nom})
        messages.success(request, f"Point de vente '{nom}' supprimé avec succès!")
        return redirect('liste_point_ventes')
    
    return render(request, 'points_vente/supprimer_point_vente.html', {
        'point_vente': point_vente
    })



# 🔒 Seulement les superusers ont accès
@login_required(login_url='login')
def admin_dashboard(request):
    from datetime import datetime, timedelta
    from django.db import models
    from django.db.models import Count, Sum, F
    
    # Vérifier que l'utilisateur est un admin
    if not _can_view_global(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    # Date du jour pour comparaison dernière connexion
    today = timezone.now().date()

    # Statistiques générales
    total_devis = Devis.objects.count()
    total_clients = Client.objects.count()
    total_produits = Produit.objects.count()

    # Statistiques des devis
    current_month = timezone.now().month
    current_year = timezone.now().year
    devis_mois_actuel = Devis.objects.filter(
        date_emission__month=current_month,
        date_emission__year=current_year
    ).count()

    # Top 5 des clients avec le plus de devis
    top_clients = Client.objects.annotate(
        nb_devis=Count('devis')
    ).order_by('-nb_devis')[:5]

    # Top commerciaux avec leurs stats
    top_commerciaux = User.objects.filter(
        profile__role='commercial'
    ).annotate(
        nb_devis=Count('devis'),  # Using the reverse relation from utilisateur field
        total_ca=Sum('devis__total_ttc')
    ).order_by('-nb_devis')[:5]

    # Données pour le graphique des devis par mois
    six_mois = timezone.now() - timedelta(days=180)
    devis_par_mois = (
        Devis.objects
        .filter(date_emission__gte=six_mois)
        .values('date_emission__month', 'date_emission__year')
        .annotate(total=Count('id'))
        .order_by('date_emission__year', 'date_emission__month')
    )

    labels = []
    data = []

    for entry in devis_par_mois:
        month = datetime(
            entry['date_emission__year'],
            entry['date_emission__month'],
            1
        ).strftime('%B %Y')
        labels.append(month)
        data.append(entry['total'])

    # Montant total des devis du mois
    montant_total_mois = Devis.objects.filter(
        date_emission__month=current_month,
        date_emission__year=current_year
    ).aggregate(total=models.Sum('total_ttc'))['total'] or 0

    context = {
        'total_devis': total_devis,
        'total_clients': total_clients,
        'total_produits': total_produits,
        'devis_mois_actuel': devis_mois_actuel,
        'montant_total_mois': montant_total_mois,
        'top_clients': top_clients,
        'top_commerciaux': top_commerciaux,
        'labels': labels,
        'data': data,
        'today': today,
    }
    
    return render(request, 'Dashboard_Admin/admin_dashboard.html', context)


@login_required(login_url='login')
def commerciaux_devis_list(request):
    """
    Affiche la liste des commerciaux (Profile.role == 'commercial')
    avec leur dernier login et la liste des devis qu'ils ont créés.
    Accessible uniquement aux superusers (admin).
    """
    # accès restreint aux admins
    if not _can_view_global(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    # Récupérer les utilisateurs ayant un profil commercial et précharger leurs devis et actions
    commerciaux = (
        User.objects
        .filter(profile__role='commercial')
        .annotate(nb_devis=Count('devis'))  # Using the reverse relation from utilisateur field
        .prefetch_related('devis', 'actioncommercial_set')
    )

    return render(request, 'commerciaux_devis_list.html', {
        'commerciaux': commerciaux,
        'is_admin_user': _is_admin_user(request.user),
    })


@login_required(login_url='login')
def create_commercial(request):
    """Permet à l'admin de créer un compte commercial et l'assigner."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')

    if request.method == 'POST':
        form = CommercialCreateForm(request.POST)
        if form.is_valid():
            try:
                user = form.save()
                messages.success(request, f"✅ Commercial {user.username} créé avec succès.")
                return redirect('liste_commerciaux')
            except Exception as e:
                messages.error(request, f"❌ Erreur lors de la création : {str(e)}")
        else:
            messages.error(request, "❌ Le formulaire contient des erreurs. Veuillez vérifier les champs.")
    else:
        form = CommercialCreateForm()

    return render(request, 'create_commercial.html', {'form': form})


@login_required(login_url='login')
def liste_commerciaux(request):
    """Liste tous les commerciaux."""
    if not _can_view_global(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    
    commerciaux = User.objects.filter(profile__role='commercial').select_related('profile', 'profile__point_vente').prefetch_related('devis')
    response = render(request, 'commercial/liste_commerciaux.html', {
        'commerciaux': commerciaux,
        'is_admin_user': _is_admin_user(request.user),
    })
    # Désactiver le cache pour éviter les problèmes d'affichage
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response


@login_required(login_url='login')
def modifier_commercial(request, user_id):
    """Permet à l'admin de modifier un commercial."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    
    user = get_object_or_404(User, id=user_id, profile__role='commercial')
    
    if request.method == 'POST':
        try:
            # Mise à jour des informations de base
            user.username = request.POST.get('username')
            user.first_name = request.POST.get('first_name', '')
            user.last_name = request.POST.get('last_name', '')
            user.email = request.POST.get('email', '')
            user.is_active = request.POST.get('is_active') == 'on'
            
            # Mise à jour du mot de passe si fourni
            new_password = request.POST.get('new_password', '').strip()
            if new_password:
                user.set_password(new_password)
            
            user.save()
            
            # Mise à jour du profil
            profile = user.profile
            point_vente_id = request.POST.get('point_vente')
            if point_vente_id:
                profile.point_vente_id = int(point_vente_id)
            else:
                profile.point_vente = None
            profile.telephone = request.POST.get('telephone', '')
            profile.save()
            
            messages.success(request, f"✅ Commercial {user.username} modifié avec succès.")
            
        except Exception as e:
            messages.error(request, f"❌ Erreur lors de la modification : {str(e)}")
            return redirect('modifier_commercial', user_id=user_id)
        
        # Redirection POST-Redirect-GET pour éviter la re-soumission du formulaire
        return redirect('liste_commerciaux')
    
    points_vente = PointVente.objects.all()
    return render(request, 'commercial/modifier_commercial.html', {
        'commercial': user,
        'points_vente': points_vente
    })


@login_required(login_url='login')
def supprimer_commercial(request, user_id):
    """Permet à l'admin de supprimer un commercial."""
    if not _is_admin_user(request.user):
        if _is_ajax_request(request):
            return JsonResponse({"success": False, "error": "Accès non autorisé."}, status=403)
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    
    user = get_object_or_404(User, id=user_id, profile__role='commercial')
    
    if request.method == 'POST':
        username = user.username
        user.delete()
        if _is_ajax_request(request):
            return JsonResponse({"success": True, "deleted": username})
        messages.success(request, f"Commercial {username} supprimé avec succès.")
        return redirect('liste_commerciaux?deleted=1')
    
    return render(request, 'commercial/supprimer_commercial.html', {'commercial': user})


@login_required(login_url='login')
def regenerate_qr_codes_view(request):
    """Permet à l'admin de régénérer tous les QR codes."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    
    from .utils import generate_qr_code
    from django.conf import settings
    
    total_devis = Devis.objects.count()
    success_count = 0
    error_count = 0
    started = False
    
    if request.method == 'POST':
        started = True
        devis_list = Devis.objects.all()
        
        for devis in devis_list:
            try:
                qr_image = generate_qr_code(devis.slug)
                if qr_image:
                    devis.qr_code.save(f"qr_{devis.slug}.png", qr_image, save=False)
                    devis.save(update_fields=['qr_code'])
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                error_count += 1
        
        messages.success(request, f"✅ {success_count} QR codes régénérés avec succès !")
        if error_count > 0:
            messages.warning(request, f"⚠️ {error_count} erreurs sont survenues.")
    
    return render(request, 'admin/regenerate_qr.html', {
        'total_devis': total_devis,
        'site_url': settings.SITE_URL,
        'started': started,
        'success_count': success_count,
        'error_count': error_count,
    })


def custom_login(request):
    """Vue de connexion avec redirection selon le rôle."""
    if request.user.is_authenticated:
        if _can_view_global(request.user):
            return redirect('admin_dashboard')
        return redirect('dashboard')

    if request.method == 'POST':
        from django.contrib.auth import authenticate, login
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            if _can_view_global(user):
                return redirect('admin_dashboard')
            return redirect('dashboard')
        else:
            messages.error(request, "Nom d'utilisateur ou mot de passe incorrect.")

    return render(request, 'Authentification/login.html')


@login_required
def liste_responsables(request):
    """Liste des responsables commerciaux — accessible admin uniquement."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    responsables = ResponsableCommercial.objects.all()
    return render(request, 'Dashboard_Admin/responsables.html', {'responsables': responsables})


@login_required
def ajouter_responsable(request):
    """Ajouter un responsable commercial."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    if request.method == 'POST':
        nom = request.POST.get('nom', '').strip()
        if nom:
            ResponsableCommercial.objects.create(nom=nom)
            messages.success(request, f"Responsable '{nom}' ajouté.")
        else:
            messages.error(request, "Le nom ne peut pas être vide.")
    return redirect('liste_responsables')


@login_required
def supprimer_responsable(request, pk):
    """Supprimer un responsable commercial."""
    if not _is_admin_user(request.user):
        messages.error(request, "Accès non autorisé.")
        return redirect('dashboard')
    responsable = get_object_or_404(ResponsableCommercial, pk=pk)
    responsable.delete()
    messages.success(request, "Responsable supprimé.")
    return redirect('liste_responsables')


def responsables_suggestions(request):
    """Retourne uniquement les commerciaux actifs pour le champ nom_responsable."""
    from django.contrib.auth.models import User
    commerciaux = User.objects.filter(
        is_active=True,
        profile__role='commercial'
    ).order_by('last_name', 'first_name')
    noms = [u.get_full_name() or u.username for u in commerciaux]
    return JsonResponse(noms, safe=False)


def client_factures(request, client_slug):
    """Portail client : affiche toutes les factures d'un client via son slug."""
    client = get_object_or_404(_scoped_clients_queryset(request.user), slug=client_slug)
    devis_list = _scoped_devis_queryset(request.user).filter(client=client).order_by('-date_emission')
    return render(request, 'clients/client_factures.html', {
        'client': client,
        'devis_list': devis_list,
    })